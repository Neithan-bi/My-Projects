import re
from docx import Document

# Load the original document
input_doc = Document(r"C:\Users\iguru\OneDrive\Desktop\Code Files\Python Files\CCP\Newpiss\Aiswaria_Final_AR_(1).docx")
full_text = "\n".join([para.text for para in input_doc.paragraphs])

# Regex pattern to extract question blocks
pattern = re.compile(
    r"Question (\d+)\n\nDifficulty Level: (.*?)\nCategory: (.*?)\n\n(.*?)(?=\n\nQuestion \d+|\Z)",
    re.DOTALL
)
matches = pattern.findall(full_text)

# Create output document
output_doc = Document()
output_doc.add_heading("AWS Questions - Formatted Table", level=1)

# Create table
table = output_doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
headers = ["S.No", "Question", "Options with Explanation", "Correct Answer", "Difficulty", "Category"]
for i, h in enumerate(headers):
    table.cell(0, i).text = h

# Helper function to extract options, correct answer, and explanation
def extract_options_explanations(block):
    options_pattern = re.compile(r"([A-D])\. (.*?)\n(✅|❌) (Correct|Incorrect)\. (.*?)\n", re.DOTALL)
    options = options_pattern.findall(block)
    option_texts = []
    correct = []

    for opt in options:
        option_line = f"{opt[0]}. {opt[1]} - {opt[3]}. {opt[4]}"
        option_texts.append(option_line)
        if opt[2] == "✅":
            correct.append(opt[0])

    return "\n".join(option_texts), ", ".join(correct)

# Fill the table with parsed data
for idx, (q_num, difficulty, category, content) in enumerate(matches, start=1):
    q_lines = content.strip().split("\n")
    question_text = q_lines[0].strip()
    options_with_expl, correct_answer = extract_options_explanations(content)

    row = table.add_row().cells
    row[0].text = str(idx)
    row[1].text = question_text
    row[2].text = options_with_expl
    row[3].text = correct_answer
    row[4].text = difficulty
    row[5].text = category

# Save the document
output_doc.save("Formatted_All_Questions.docx")
